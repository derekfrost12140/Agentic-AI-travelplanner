import streamlit as st
import os
from datetime import datetime, timedelta
from travel_agent import TravelAgent
from langchain.schema import HumanMessage, AIMessage
import time
import json
import io

# Optional import for DOCX support
try:
    from docx import Document
    from docx.shared import Inches
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Page configuration
st.set_page_config(
    page_title="AI Travel Agent",
    page_icon="✈️",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Set dark theme
st.markdown("""
<style>
    .stApp {
        background: linear-gradient(135deg, #1e3c72 0%, #2a5298 100%);
        color: white;
    }
</style>
""", unsafe_allow_html=True)

# Custom CSS for better styling
st.markdown("""
<style>
    /* Main headers */
    .main-header {
        font-size: 3rem;
        font-weight: bold;
        text-align: center;
        color: #00d4aa;
        margin-bottom: 2rem;
        text-shadow: 0 2px 4px rgba(0,0,0,0.3);
    }
    .sub-header {
        font-size: 1.5rem;
        color: #ffffff;
        text-align: center;
        margin-bottom: 2rem;
        opacity: 0.9;
    }
    
    /* Chat messages with better contrast */
    .chat-message {
        padding: 1.5rem;
        border-radius: 1rem;
        margin-bottom: 1.5rem;
        border-left: 4px solid;
        box-shadow: 0 4px 12px rgba(0,0,0,0.15);
        font-size: 1rem;
        line-height: 1.6;
    }
    .user-message {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        border-left-color: #4c63d2;
        color: white;
        text-shadow: 0 1px 2px rgba(0,0,0,0.1);
    }
    .assistant-message {
        background: linear-gradient(135deg, #f093fb 0%, #f5576c 100%);
        border-left-color: #e91e63;
        color: white;
        text-shadow: 0 1px 2px rgba(0,0,0,0.1);
    }
    
    /* Buttons with better styling */
    .stButton > button {
        width: 100%;
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        color: white;
        border: none;
        padding: 0.75rem 1.5rem;
        border-radius: 0.75rem;
        font-weight: bold;
        font-size: 1rem;
        transition: all 0.3s ease;
        box-shadow: 0 4px 12px rgba(102, 126, 234, 0.4);
    }
    .stButton > button:hover {
        background: linear-gradient(135deg, #5a6fd8 0%, #6a4190 100%);
        transform: translateY(-2px);
        box-shadow: 0 6px 20px rgba(102, 126, 234, 0.6);
    }
    
    /* Sidebar styling */
    .sidebar .sidebar-content {
        background: linear-gradient(180deg, #2c3e50 0%, #34495e 100%);
        border-right: 2px solid #34495e;
    }
    
    /* Text input styling */
    .stTextInput > div > div > input {
        background-color: #34495e;
        color: white;
        border: 2px solid #4a90e2;
        border-radius: 0.5rem;
        padding: 0.75rem;
        font-size: 1rem;
    }
    .stTextInput > div > div > input:focus {
        border-color: #00d4aa;
        box-shadow: 0 0 0 2px rgba(0, 212, 170, 0.2);
    }
    
    /* Text area styling */
    .stTextArea > div > div > textarea {
        background-color: #34495e;
        color: white;
        border: 2px solid #4a90e2;
        border-radius: 0.75rem;
        padding: 1rem;
        font-size: 1rem;
        line-height: 1.5;
    }
    .stTextArea > div > div > textarea:focus {
        border-color: #00d4aa;
        box-shadow: 0 0 0 2px rgba(0, 212, 170, 0.2);
    }
    
    /* Form submit button */
    .stForm > div > div > button {
        background: linear-gradient(135deg, #00d4aa 0%, #00b894 100%);
        color: white;
        border: none;
        padding: 0.75rem 2rem;
        border-radius: 0.75rem;
        font-weight: bold;
        font-size: 1.1rem;
        transition: all 0.3s ease;
        box-shadow: 0 4px 12px rgba(0, 212, 170, 0.4);
    }
    .stForm > div > div > button:hover {
        background: linear-gradient(135deg, #00b894 0%, #00a085 100%);
        transform: translateY(-2px);
        box-shadow: 0 6px 20px rgba(0, 212, 170, 0.6);
    }
    
    /* Info boxes */
    .stAlert {
        background: linear-gradient(135deg, #74b9ff 0%, #0984e3 100%);
        color: white;
        border-radius: 0.75rem;
        border: none;
        padding: 1rem;
    }
    
    /* Headers in sidebar */
    .sidebar h1, .sidebar h2, .sidebar h3 {
        color: #00d4aa;
        text-shadow: 0 1px 2px rgba(0,0,0,0.3);
    }
    
    /* Divider styling */
    hr {
        border: none;
        height: 2px;
        background: linear-gradient(90deg, transparent, #00d4aa, transparent);
        margin: 2rem 0;
    }
    
    /* Better spacing */
    .block-container {
        padding-top: 2rem;
        padding-bottom: 2rem;
    }
    
    /* Download buttons styling */
    .download-section {
        background: rgba(255,255,255,0.05);
        border-radius: 1rem;
        padding: 1rem;
        margin: 1rem 0;
        border: 1px solid rgba(255,255,255,0.1);
    }
    
    /* Smaller, more beautiful download buttons */
    .download-button {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        color: white;
        border: none;
        padding: 0.5rem 1rem;
        border-radius: 0.5rem;
        font-size: 0.9rem;
        font-weight: bold;
        transition: all 0.3s ease;
        box-shadow: 0 2px 8px rgba(102, 126, 234, 0.3);
        margin: 0.25rem;
    }
    
    .download-button:hover {
        background: linear-gradient(135deg, #5a6fd8 0%, #6a4190 100%);
        transform: translateY(-1px);
        box-shadow: 0 4px 12px rgba(102, 126, 234, 0.5);
    }
    
    /* Compact download section */
    .compact-download {
        margin: 0.5rem 0;
        padding: 0.5rem;
    }
</style>
""", unsafe_allow_html=True)

def initialize_session_state():
    """Initialize session state variables"""
    if 'messages' not in st.session_state:
        st.session_state.messages = []
    if 'agent' not in st.session_state:
        st.session_state.agent = None
    if 'chat_history' not in st.session_state:
        st.session_state.chat_history = []
    if 'planning_step' not in st.session_state:
        st.session_state.planning_step = None
    if 'flight_options' not in st.session_state:
        st.session_state.flight_options = []
    if 'selected_flight' not in st.session_state:
        st.session_state.selected_flight = None
    if 'hotel_options' not in st.session_state:
        st.session_state.hotel_options = []
    if 'selected_hotel' not in st.session_state:
        st.session_state.selected_hotel = None

def create_travel_agent():
    """Create and return a travel agent instance"""
    try:
        return TravelAgent()
    except ValueError as e:
        st.error(f"Error initializing travel agent: {e}")
        st.info("Please make sure you have set your OPENAI_API_KEY in the .env file")
        return None

def display_chat_message(message, is_user=False):
    """Display a chat message with proper styling"""
    if is_user:
        st.markdown(f"""
        <div class="chat-message user-message">
            <strong>You:</strong><br>
            {message}
        </div>
        """, unsafe_allow_html=True)
    else:
        st.markdown(f"""
        <div class="chat-message assistant-message">
            <strong>AI Travel Agent:</strong><br>
            {message}
        </div>
        """, unsafe_allow_html=True)

def create_itinerary_document(conversation_history, format_type="txt"):
    """Create an itinerary document from conversation history"""
    
    # Extract travel planning information from conversation
    itinerary_content = "✈️ AI Travel Agent - Travel Itinerary\n"
    itinerary_content += "=" * 50 + "\n\n"
    itinerary_content += f"Generated on: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}\n\n"
    
    # Add conversation summary
    itinerary_content += "📋 CONVERSATION SUMMARY\n"
    itinerary_content += "-" * 30 + "\n"
    
    for message in conversation_history:
        if message["role"] == "user":
            itinerary_content += f"👤 You: {message['content']}\n\n"
        else:
            itinerary_content += f"🤖 AI Travel Agent: {message['content']}\n\n"
    
    itinerary_content += "\n" + "=" * 50 + "\n"
    itinerary_content += "Built with 🔥 passion by Devansh Swami\n"
    itinerary_content += "AI Travel Agent v1.0\n"
    
    if format_type == "txt":
        return itinerary_content
    elif format_type == "json":
        return json.dumps({
            "title": "AI Travel Agent - Travel Itinerary",
            "generated_on": datetime.now().isoformat(),
            "conversation": conversation_history,
            "author": "Devansh Swami",
            "version": "1.0"
        }, indent=2)
    elif format_type == "docx":
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx package not available")
        
        # Create Word document
        doc = Document()
        
        # Title
        title = doc.add_heading('✈️ AI Travel Agent - Travel Itinerary', 0)
        title.alignment = 1  # Center alignment
        
        # Date
        date_para = doc.add_paragraph(f"Generated on: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}")
        date_para.alignment = 1
        
        doc.add_paragraph()  # Spacing
        
        # Conversation
        doc.add_heading('📋 CONVERSATION SUMMARY', level=1)
        
        for message in conversation_history:
            if message["role"] == "user":
                p = doc.add_paragraph()
                p.add_run("👤 You: ").bold = True
                p.add_run(message['content'])
            else:
                p = doc.add_paragraph()
                p.add_run("🤖 AI Travel Agent: ").bold = True
                p.add_run(message['content'])
            doc.add_paragraph()  # Spacing
        
        # Footer
        doc.add_paragraph("=" * 50)
        footer = doc.add_paragraph("Built with 🔥 passion by Devansh Swami")
        footer.alignment = 1
        version = doc.add_paragraph("AI Travel Agent v1.0")
        version.alignment = 1
        
        # Save to bytes
        doc_bytes = io.BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)
        return doc_bytes.getvalue()
    
    return itinerary_content

def download_itinerary_button(conversation_history):
    """Create download buttons for different formats"""
    
    if not conversation_history:
        return
    
    # Smaller, more beautiful download section
    st.markdown("""
    <div class="download-section">
        <h4 style="margin-bottom: 0.5rem; color: #00d4aa;">📥 Download Itinerary</h4>
    """, unsafe_allow_html=True)
    
    # Create a compact container for download buttons
    col1, col2, col3 = st.columns(3)
    
    with col1:
        # Download as TXT
        txt_content = create_itinerary_document(conversation_history, "txt")
        st.download_button(
            label="📄 TXT",
            data=txt_content,
            file_name=f"travel_itinerary_{datetime.now().strftime('%Y%m%d_%H%M')}.txt",
            mime="text/plain",
            use_container_width=True,
            help="Download as simple text file"
        )
    
    with col2:
        # Download as JSON
        json_content = create_itinerary_document(conversation_history, "json")
        st.download_button(
            label="📊 JSON",
            data=json_content,
            file_name=f"travel_itinerary_{datetime.now().strftime('%Y%m%d_%H%M')}.json",
            mime="application/json",
            use_container_width=True,
            help="Download as structured JSON data"
        )
    
    with col3:
        # Download as DOCX
        if DOCX_AVAILABLE:
            try:
                docx_content = create_itinerary_document(conversation_history, "docx")
                st.download_button(
                    label="📝 DOCX",
                    data=docx_content,
                    file_name=f"travel_itinerary_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                    help="Download as professional Word document"
                )
            except Exception as e:
                st.error(f"Error creating DOCX: {str(e)}")
        else:
            st.info("📝 Install python-docx for DOCX export")
            st.markdown("```bash\npip install python-docx\n```")
    
    # Close the download section div
    st.markdown("</div>", unsafe_allow_html=True)

def main():
    # Initialize session state
    initialize_session_state()
    
    # Load API key from .env file
    from dotenv import load_dotenv
    load_dotenv()
    
    # Header
    st.markdown('<h1 class="main-header">✈️ AI Travel Agent</h1>', unsafe_allow_html=True)
    st.markdown('<p class="sub-header">Plan your perfect trip with AI assistance</p>', unsafe_allow_html=True)
    
    # Sidebar for configuration
    with st.sidebar:
        st.header("🔧 Configuration")
        
        # Check if API key exists in environment
        env_api_key = os.getenv("OPENAI_API_KEY")
        
        if env_api_key and env_api_key != "your_openai_api_key_here":
            st.success("✅ API key loaded from .env file!")
            
            # Initialize agent if not already done
            if st.session_state.agent is None:
                with st.spinner("Initializing AI Travel Agent..."):
                    st.session_state.agent = create_travel_agent()
        else:
            # API Key input with better styling
            st.markdown("### 🔑 API Configuration")
            api_key = st.text_input(
                "OpenAI API Key",
                type="password",
                help="Enter your OpenAI API key to use the travel agent",
                placeholder="sk-..."
            )
            
            if not api_key:
                st.info("💡 **Get your API key from:** [OpenAI Platform](https://platform.openai.com/api-keys)")
                st.info("💡 **Or add it to your .env file:** `OPENAI_API_KEY=your_key_here`")
            else:
                st.success("✅ API key configured!")
                os.environ["OPENAI_API_KEY"] = api_key
                
                # Initialize agent if not already done
                if st.session_state.agent is None:
                    with st.spinner("Initializing AI Travel Agent..."):
                        st.session_state.agent = create_travel_agent()
        
        st.markdown("---")
        
        # Quick actions
        st.header("🚀 Quick Actions")
        
        if st.button("Plan a Trip to Paris"):
            if st.session_state.agent:
                user_message = "I want to plan a 5-day trip to Paris next month. Can you help me with flights, hotels, and activities?"
                st.session_state.messages.append({"role": "user", "content": user_message})
                
                # Get AI response
                with st.spinner("AI Travel Agent is thinking..."):
                    try:
                        response = st.session_state.agent.chat(user_message, st.session_state.chat_history)
                        st.session_state.messages.append({"role": "assistant", "content": response})
                        st.session_state.chat_history.extend([
                            HumanMessage(content=user_message),
                            AIMessage(content=response)
                        ])
                    except Exception as e:
                        error_msg = f"Sorry, I encountered an error: {str(e)}. Please try again."
                        st.session_state.messages.append({"role": "assistant", "content": error_msg})
                st.rerun()
            else:
                st.error("Please configure your OpenAI API key first")
        
        if st.button("Find Flights to Tokyo"):
            if st.session_state.agent:
                user_message = "I need to find flights from New York to Tokyo for next week. What are my options?"
                st.session_state.messages.append({"role": "user", "content": user_message})
                
                # Get AI response
                with st.spinner("AI Travel Agent is thinking..."):
                    try:
                        response = st.session_state.agent.chat(user_message, st.session_state.chat_history)
                        st.session_state.messages.append({"role": "assistant", "content": response})
                        st.session_state.chat_history.extend([
                            HumanMessage(content=user_message),
                            AIMessage(content=response)
                        ])
                    except Exception as e:
                        error_msg = f"Sorry, I encountered an error: {str(e)}. Please try again."
                        st.session_state.messages.append({"role": "assistant", "content": error_msg})
                st.rerun()
            else:
                st.error("Please configure your OpenAI API key first")
        
        if st.button("Get Weather for New York"):
            if st.session_state.agent:
                user_message = "What's the weather like in New York this weekend?"
                st.session_state.messages.append({"role": "user", "content": user_message})
                
                # Get AI response
                with st.spinner("AI Travel Agent is thinking..."):
                    try:
                        response = st.session_state.agent.chat(user_message, st.session_state.chat_history)
                        st.session_state.messages.append({"role": "assistant", "content": response})
                        st.session_state.chat_history.extend([
                            HumanMessage(content=user_message),
                            AIMessage(content=response)
                        ])
                    except Exception as e:
                        error_msg = f"Sorry, I encountered an error: {str(e)}. Please try again."
                        st.session_state.messages.append({"role": "assistant", "content": error_msg})
                st.rerun()
            else:
                st.error("Please configure your OpenAI API key first")
        
        st.markdown("---")
        
                # Clear chat button
        if st.button("🗑️ Clear Chat"):
            st.session_state.messages = []
            st.session_state.chat_history = []
            st.rerun()
    
    # Main chat area
    col1, col2, col3 = st.columns([1, 2, 1])
    
    with col2:
        # Display chat messages
        for idx, message in enumerate(st.session_state.messages):
            if message["role"] == "user":
                display_chat_message(message["content"], is_user=True)
            else:
                display_chat_message(message["content"], is_user=False)


        
        # Chat input with better styling
        if st.session_state.agent:
            st.success("🤖 AI Travel Agent is ready! Start chatting below.")
            
            st.markdown("### 💬 Chat with AI Travel Agent")
            with st.form("chat_form", clear_on_submit=True):
                user_input = st.text_area(
                    "Ask me about travel planning, flights, hotels, weather, or recommendations!",
                    height=120,
                    placeholder="e.g., I want to plan a trip to Paris for next month...",
                    help="Type your travel question here and press Send"
                )
                
                col1, col2, col3 = st.columns([1, 2, 1])
                with col2:
                    submit_button = st.form_submit_button("🚀 Send Message", use_container_width=True)
                
                if submit_button and user_input.strip():
                    # Add user message to chat
                    st.session_state.messages.append({"role": "user", "content": user_input})
                    
                    # Get AI response
                    with st.spinner("AI Travel Agent is thinking..."):
                        try:
                            response = st.session_state.agent.chat(
                                user_input, 
                                st.session_state.chat_history
                            )
                            
                            # Debug: Check if response is empty or too short
                            if not response or len(response.strip()) < 10:
                                response = "I apologize, but I didn't receive a proper response. Please try asking your question again."
                            
                            # Add AI response to chat
                            st.session_state.messages.append({"role": "assistant", "content": response})
                            
                            # Update chat history for LangChain
                            st.session_state.chat_history.extend([
                                HumanMessage(content=user_input),
                                AIMessage(content=response)
                            ])
                            
                        except Exception as e:
                            error_msg = f"Sorry, I encountered an error: {str(e)}. Please try again."
                            st.session_state.messages.append({"role": "assistant", "content": error_msg})
                    
                    st.rerun()
            
            # Download section (only show if there are messages) - positioned below chat
            if st.session_state.messages:
                st.markdown("---")
                download_itinerary_button(st.session_state.messages)
                
        elif env_api_key and env_api_key != "your_openai_api_key_here":
            st.info("🔄 Initializing AI Travel Agent... Please wait.")
        else:
            st.warning("⚠️ Please configure your OpenAI API key to start chatting.")
            
            # Show example conversation with better styling
            st.markdown("### 💡 Example Conversation")
            
            # Create example conversation using Streamlit components instead of raw HTML
            st.markdown("**👤 You:** I want to plan a trip to Paris for next month. Can you help me?")
            
            st.markdown("**🤖 AI Travel Agent:** I'd be happy to help you plan your trip to Paris! To provide the best recommendations, I need a few details:")
            
            st.markdown("""
            - What dates are you planning to travel?
            - How many people are traveling?
            - What's your budget range?
            - What are your main interests (culture, food, shopping, etc.)?
            - Where will you be departing from?
            """)
            
            st.markdown("Once I have these details, I can help you find flights, hotels, and create a personalized itinerary!")
    
    # Footer with better styling
    st.markdown("---")
    st.markdown("""
    <div style='text-align: center; color: #ffffff; opacity: 0.8; padding: 2rem 0;'>
        <p style='font-size: 1.1rem; margin-bottom: 0.5rem;'>Built with 🔥 passion using LangChain and Streamlit</p>
        <p style='font-size: 0.9rem;'>This is a proof of concept demonstrating agentic AI capabilities</p>
        <p style='font-size: 1rem; margin-top: 0.5rem; color: #00d4aa; font-weight: bold;'>by Devansh Swami</p>
        <div style='margin-top: 1rem;'>
            <span style='background: linear-gradient(135deg, #00d4aa, #00b894); padding: 0.5rem 1rem; border-radius: 2rem; font-size: 0.8rem;'>AI Travel Agent v1.0</span>
        </div>
    </div>
    """, unsafe_allow_html=True)

if __name__ == "__main__":
    main() 